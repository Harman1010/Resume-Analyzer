import os

from docx import Document

from docx.shared import Pt

from schemas.rewriter_schema import RewriteSchema


def get_resume(rewriten_resume : RewriteSchema,output_path: str = "output/ATS_Optimized_Resume.docx")->str:

    document = Document()

    title = document.add_heading("AI Optimized Resume",level=1)

    title.runs[0].font.size = Pt(20)

    document.add_heading("Profile",level=2)

    document.add_paragraph(rewriten_resume.profile)

    document.add_heading("Skills",level=2)

    for skill in rewriten_resume.skills:
        document.add_paragraph(skill,style="List Bullet")

    document.add_heading("Projects",level=2)

    for project in rewriten_resume.projects:
        document.add_paragraph(project,style="List Bullet")

    document.add_heading("Experience",level=2)

    for experience in rewriten_resume.experience:
        document.add_paragraph(experience,style="List Bullet")

    document.add_heading("Education",level=2)

    for education in rewriten_resume.education:
        document.add_paragraph(education,style="List Bullet")

    document.add_heading("Achievements",level=2)

    for achievements in rewriten_resume.achievements:
        document.add_paragraph(achievements,style="List Bullet")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    document.save(output_path)

    return output_path

    

#Input - Rewrite
#Output - Python-docx